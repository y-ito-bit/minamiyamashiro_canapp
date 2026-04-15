from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

os.chdir(os.path.dirname(os.path.abspath(__file__)))

doc = Document()

# Page setup - A4, narrow margins for 2 pages
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

# Default font - Japanese
style = doc.styles['Normal']
style.font.name = 'Yu Gothic'
style.font.size = Pt(10)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts')
    rpr.append(rfonts)
rfonts.set(qn('w:eastAsia'), 'Yu Gothic')
rfonts.set(qn('w:ascii'), 'Yu Gothic')
rfonts.set(qn('w:hAnsi'), 'Yu Gothic')

def set_jp_font(run, size=10, bold=False, color=None):
    run.font.name = 'Yu Gothic'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:eastAsia'), 'Yu Gothic')
    rfonts.set(qn('w:ascii'), 'Yu Gothic')
    rfonts.set(qn('w:hAnsi'), 'Yu Gothic')
    rpr.append(rfonts)

def add_heading(text, size=13, color=RGBColor(0x1F, 0x4E, 0x79)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_jp_font(r, size=size, bold=True, color=color)
    return p

def add_para(text, size=10, bold=False, color=None, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_jp_font(r, size=size, bold=bold, color=color)
    return p

def add_bullet(text, size=9.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_jp_font(r, size=size)
    return p

# === Title ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
r = title.add_run('南山城学園 5年目研修アプリ ご利用マニュアル')
set_jp_font(r, size=16, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(6)
r = sub.add_run('〜 AI先輩との対話で「あなたの強み」を再発見 〜')
set_jp_font(r, size=10, color=RGBColor(0x55, 0x55, 0x55))

# Intro
add_para('本アプリは、5年間の業務を振り返り、ご自身の「強み」や「専門性」を再発見するためのキャリア支援ツールです。AIキャリアコーチとの対話を通じて、あなただけの「お仕事ポートフォリオ」を作成いただけます。', size=10, space_after=4)

# === Step 1 ===
add_heading('STEP 1  プロフィール・性格タイプの入力', size=12)
add_para('お名前（フルネーム／ニックネーム可）を入力し、MBTI性格診断の質問に直感でお答えください。AIがあなたに最適なアプローチで質問するための準備となります。', size=9.5, space_after=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
p.add_run().add_picture('フェーズ１名前入力.jpg', width=Cm(8.5))

# === Step 2 ===
add_heading('STEP 2  AI先輩とのキャリア面談（チャット）', size=12)
add_para('「過去5年間でどのような業務をしてきましたか？」など、優しい先輩目線の質問にチャット感覚でお答えください。成功体験はもちろん、失敗から学んだ経験の中にも強みは隠れています。', size=9.5, space_after=2)
add_para('▶ マイクボタンで音声入力にも対応（初回のみブラウザのマイク許可が必要）。', size=9.5, bold=True, color=RGBColor(0xC0, 0x39, 0x2B), space_after=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
p.add_run().add_picture('フェーズ２チャット欄.jpg', width=Cm(8.5))

# === Step 3 ===
add_heading('STEP 3  「お仕事ポートフォリオ」の確認', size=12)
add_para('対話が十分に進むと、画面に「✨診断レポートを見る」ボタンが表示されます。タップすると、あなたの5年間の価値が詰まったレポートが完成します。', size=9.5, space_after=2)

# Two images side by side using a table
table = doc.add_table(rows=1, cols=2)
table.autofit = True
for cell in table.rows[0].cells:
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
table.rows[0].cells[0].paragraphs[0].add_run().add_picture('フェーズ３強みレポート１.jpg', width=Cm(8.0))
table.rows[0].cells[1].paragraphs[0].add_run().add_picture('フェーズ３強みレポート２.jpg', width=Cm(8.0))

add_para('●強みマップ：「実務・専門 / 連携 / 共感 / しなやかさ」の4指標で強みを可視化   ●見つかった専門性と強み：日々の業務に潜む「確かな価値」を言語化   ●強みからのメッセージ：今後のキャリアパスへのヒント', size=9, space_after=4)

# === 注意点 (bottom) ===
# Add a colored box for the warning section using a single-cell table with shading
def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

warn_table = doc.add_table(rows=1, cols=1)
warn_table.autofit = True
cell = warn_table.rows[0].cells[0]
shade_cell(cell, 'FFF4E5')

# Set borders
tc_pr = cell._tc.get_or_add_tcPr()
tc_borders = OxmlElement('w:tcBorders')
for b in ['top','left','bottom','right']:
    el = OxmlElement(f'w:{b}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), '8')
    el.set(qn('w:color'), 'E67E22')
    tc_borders.append(el)
tc_pr.append(tc_borders)

cp = cell.paragraphs[0]
cp.paragraph_format.space_after = Pt(2)
r = cp.add_run('⚠ ご利用上の注意点（必ずお読みください）')
set_jp_font(r, size=11, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))

def cell_para(cell, text, size=9, bold=False, color=None):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_jp_font(r, size=size, bold=bold, color=color)
    return p

cell_para(cell, '■ 個人情報の入力禁止', size=9.5, bold=True)
cell_para(cell, '対話では具体的なエピソードを伺いますが、利用者様の実名・住所・病歴など個人を特定できる情報は絶対に入力しないでください。必ず匿名・伏せ字でご記入ください。')
cell_para(cell, '  ⭕ 良い例：「Aさん」「80代の女性の利用者様」    ❌ 悪い例：「山田太郎さん」「京都府〇〇町の〜」', size=9)

cell_para(cell, '■ AI学習設定について（ガバナンス上の推奨）', size=9.5, bold=True)
cell_para(cell, '本アプリでは入力内容がAIの機械学習に利用されない設定としておりますが、情報ガバナンスの観点から、機微な情報・機密情報・個人情報の入力はお控えいただくことを推奨いたします。技術設定に依存せず「入力しない」ことが最も確実な情報保護となります。')

cell_para(cell, '■ 生成が途中で止まった場合の対処', size=9.5, bold=True)
cell_para(cell, '稀にチャットのやりとり中にAIの生成が途中で止まることがあります。その際は入力欄に「続けて」とお送りいただくと、途中から生成を再開しますのでご安心ください。')

cell_para(cell, '■ 理解度ゲージの目安', size=9.5, bold=True)
cell_para(cell, '対話画面の理解度ゲージは80%程度まで進めば十分です。100%を目指す必要はなく、80%前後で「✨診断レポートを見る」ボタンからレポートを作成いただけます。')

cell_para(cell, '■ 業務利用のお願い', size=9.5, bold=True)
cell_para(cell, '本アプリは業務目的でご利用ください。生成された「お仕事ポートフォリオ」は、今後の目標設定やキャリアステップアップにご活用いただければ幸いです。')

out_path = 'ご利用マニュアル.docx'
doc.save(out_path)
print('saved:', out_path)
